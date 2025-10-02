#!/usr/bin/env python3
"""
Convert Markdown files to DOCX for Maszynownia client delivery
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def md_to_docx(md_file, docx_file):
    """Convert markdown file to formatted DOCX"""

    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Parse and add content
    lines = content.split('\n')

    for line in lines:
        line = line.rstrip()

        # Skip empty lines at start
        if not line and len(doc.paragraphs) == 0:
            continue

        # H1 - Main title
        if line.startswith('# '):
            text = line.replace('# ', '')
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H2 - Section headings
        elif line.startswith('## '):
            text = line.replace('## ', '')
            p = doc.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H3 - Subsection
        elif line.startswith('### '):
            text = line.replace('### ', '')
            p = doc.add_heading(text, level=3)

        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph()

        # Bold text patterns
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.bold = True
                else:
                    p.add_run(part)

        # Bullet points
        elif line.startswith('- '):
            text = line.replace('- ', '')
            p = doc.add_paragraph(text, style='List Bullet')

        # Empty line
        elif not line:
            doc.add_paragraph()

        # Regular paragraph
        else:
            doc.add_paragraph(line)

    # Save document
    doc.save(docx_file)
    print(f"✅ Created: {docx_file}")

# Files to convert
files = [
    ('FAQ_EMS.md', 'FAQ_EMS.docx'),
    ('FAQ_STREFA_GIMNASTYKI.md', 'FAQ_STREFA_GIMNASTYKI.docx'),
    ('BLOG_EMS_PRZYKLADOWY_POST.md', 'BLOG_EMS_PRZYKLADOWY_POST.docx'),
    ('INSTRUKCJA_GOOGLE_SEARCH_CONSOLE.md', 'INSTRUKCJA_GOOGLE_SEARCH_CONSOLE.docx'),
    ('INSTRUKCJA_CMS.md', 'INSTRUKCJA_CMS.docx'),
]

print("🔄 Converting Markdown to DOCX...\n")

for md_file, docx_file in files:
    try:
        md_to_docx(md_file, docx_file)
    except FileNotFoundError:
        print(f"⚠️  Skipped: {md_file} (not found)")
    except Exception as e:
        print(f"❌ Error with {md_file}: {e}")

print("\n✅ Conversion complete!")
print("\n📁 DOCX files ready for client delivery:")
for _, docx_file in files:
    print(f"   • {docx_file}")
